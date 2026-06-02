"""
- 워드 파일의 첫 번째 테이블(월간 리뷰 섹션)만 추출하여 사용
- pip install openai python-docx
"""

import os
import re
import openai
from docx import Document

# ════════════════════════════════════════════════════════════════
# config.py에서 설정 import
# ════════════════════════════════════════════════════════════════

try:
    from config import REVIEW_CONFIG
    OPENAI_API_KEY = REVIEW_CONFIG['openai_api_key']
    MODEL = REVIEW_CONFIG['model']
    TEMPERATURE = REVIEW_CONFIG['temperature']
    MAX_TOKENS = REVIEW_CONFIG['max_tokens']
    TRAINING_DOCX_DIR = REVIEW_CONFIG['training_dir']
    TEST_DOCX_PATH = REVIEW_CONFIG['test_file']
    OUTPUT_DIR = REVIEW_CONFIG['output_dir']
except ImportError:
    # config.py가 없을 경우 기본값 사용
    OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
    MODEL = "gpt-4o"
    TEMPERATURE = 0.1
    MAX_TOKENS = 4000
    TRAINING_DOCX_DIR = r"C:\Users\intern9\truston_quant_dev\pdf_train"
    TEST_DOCX_PATH = r"C:\Users\intern9\truston_quant_dev\pdf_test\트러스톤-2026년 4월 주식 전망.docx"
    OUTPUT_DIR = r"C:\Users\intern9\truston_quant_dev\pdf_output"

client = openai.OpenAI(api_key=OPENAI_API_KEY)


# 시스템 프롬프트

SYSTEM_PROMPT = """당신은 국내 증권사 월간 주식시장 리뷰 보고서를 읽고 요약 보고서를 작성하는 전문 애널리스트입니다.

[형식-매우 중요]
- "[N월 리뷰]" 로 시작, 다음 줄에 "1. 전월동향: KOSPI XXX.XXpt(+X.XX%)" 형식의 헤더 작성.
- 그 다음 줄부터 "- " 를 붙이고 본문 시작.
- 본문은 줄바꿈 없이 하나의 문단. 문장 사이는 마침표 + 공백.
- 제목·머리말·번호 매기기 금지 (위 헤더 항목 제외).
- 원문의 형식을 최대한 그대로 유지하며 [금지]에 해당 하지 않는 내용은 최대한 누락하지 않고 작성한다.

[본문 구성 순서 ]
1. 월간 KOSPI 개요: 등락률, 지수 마감치, 주요 상승/하락 원인 요약.
2. 핵심 이벤트별 동향: 원문에 언급된 주요 이슈(경제지표, 기업실적, 정책 이벤트, 지정학 리스크 등)를 각각 자세하게 서술.
3. 업종별 동향: 강세 업종과 약세 업종, 수급 특이사항 포함.
4. KOSDAQ, KOSPI, KOSPI200 동향: 지수 마감치와 등락률, 업종별 특이사항.

[수치 규칙]
- 지수, 등락률(%), 거래대금, 수급 금액 등 원문에 언급된 모든 수치를 누락 없이 기재.
- 이벤트에 붙은 괄호 수치(예: "+6.28%", "3,386.05p", "+16.3%")는 반드시 포함.
- 연속 증감 횟수, 역대 순위 등이 언급된 경우 그대로 포함.

[문체]
- 종결: "~하였음", "~기록하였음", "~보였음", "~증가", "~감소" 등.
- 연결어: "한편,", "특히,", "아울러", "다만," 등 자연스럽게 사용.
- 증가: "+X.X%", 감소: "-X.X%".

[금지 - 가장 중요]
- 본문 내부 줄바꿈 금지.
- 수치 누락 금지(괄호 내 수치 포함).
- 한국과 관련 없는 글로벌 매크로 관련 내용 포함 금지
"""

# Training 보고서 (N월 전망 docx → N-1월 리뷰 요약)

EXAMPLE_REPORTS = {

    "(대신) 2025년 6월 주식시장 전망.docx": {
        "target_month": "2025년 5월",
        "next_month"  : "2025년 6월",
        "report": """[5월 리뷰]
1.	전월동향: KOSPI 2,697.67pt(+5.52%)
- 5월 KOSPI는 미국의 관세정책이 협상 국면으로 접어들며 4월 중순 이후 반등 추세 지속. 특히 미-중간 관세조치 유예와 협상 프로토콜 구축, EU와의 협상 진행 등 관세정책 불확실성으로 인한 경기&물가 우려 크게 낮아졌다는 평가. Moody's의 미국 신용등급 하락과 미국 장기채 수급 이슈로 한때 국채금리 변동성 증가했으나, 금리 상승 둔화, 국내 대선이후 코리아 디스카운트 해소 기대감 등 반영되며 상승추세 지속. 대선을 앞두고 여당과 야당이 모두 자본시장 선진화 및 한국증시 저평가 해소를 정책 방향으로 제시. 민주당은 상법개정 및 자사주 강제소각, 국민의 힘은 배당소득 분리과세 및 세제혜택 등을 공약. 코리아 밸류업지수 리밸런싱 등 우호적 이슈들이 지속되면서 저평가 종목들을 중심으로 밸류 정상화 기대감 유입. 업종별로도 정책 수혜 기대 업종 외에도 화장품, 금융, 유통 등으로 주도주 흐름이 확산되며 순환매 흐름이 나타났고, 실적 모멘텀을 동반한 대형 가치주 중심으로 낙폭과대 업종 전반에 매수세가 유입. 반면, AI 과잉 투자에 대한 경계감과 미국 기술주 조정 여파로 반도체 업종은 상대적으로 약세. 외국인은 10개월 만에 월간 기준 순매수로 전환하며 수급 환경도 뚜렷하게 개선되는 가운데 KOSPI 지수는 2,697.67pt로 마감. 스타일은 대형주 비중이 높은 KOSPI200 지수는 359.62pt로 +6.16% 상승하였고 중소형주 비중이 높은 KOSDAQ 지수는 +2.39% 상승하며 마감.""",
    },

    "(대신) 트러스톤-2025년 5월 주식 전망.docx": {
        "target_month": "2025년 4월",
        "next_month"  : "2025년 5월",
        "report": """[4월 리뷰]
1.	전월동향: KOSPI 2,556.61pt(+3.04%)
- 4월 KOSPI는 월초 상호관세 발표 이후 확산된 미국 신뢰도 하락, 불확실성이 위험자산 회피로 이어졌음. 상승밴드 하단을 이탈한 뒤 주요 지지선들이 차례로 붕괴되며 큰 폭으로 하락. 이후 트럼프가 상호관세 90일 유예조치 등으로 한발 물러선 뒤 글로벌 증시 반등과 함께 V자 반등에 성공. 국내 정치적 불확실성 완화 등 더해지며 원달러 환율 1,400원대 초반으로 하락 안정, 대선국면 전환으로 인한 정책 기대, 미국의 주요 동맹국으로 빠르게 통상협의에 돌입함에 따른 관세완화 기대감, 중국의 경기부양 등 더해지면서 KOSPI 상승 전환. 4월 4일 윤석열 대통령의 탄핵심판 선고. 12월 14일 탄핵소추 이후 111일만, 2월 25일 변론 종결 이후 38일만. 100일 넘게 이어진 행정부 수장 공백이 해결로 정치적 공백 안정, 원달러 환율 안정. 대선국면 전환 이후 주요 후보들의 공약에 따라 수혜 기대감 등 유입되는 가운데 정부와 여야는 추가경정 예산안 추진이 힘을 받는 중. 방산·조선 업종은 유럽의 재무장 수요 확대, 대규모 수출 기대감, 원화 강세에 따른 비용 부담 완화 등 복합적인 호재에 힘입어 외국인 매수세가 유입되며 강세를 보이며 2,556.61pt로 마감. KOSPI 지수는 +3.04% 상승하며 2,556.61pt로 마감. KOSPI200 지수는 338.74pt로 +1.91% 상승. KOSDAQ 지수는 +6.60% 상승하며 717.24pt로 마감.
""",
    },

    "(대신) 트러스톤-2025년 7월 주식 전망.docx": {
        "target_month": "2025년 6월",
        "next_month"  : "2025년 7월",
        "report": """[6월 리뷰]
1.	전월동향: KOSPI 3,071.70pt(+13.86%)
- 6월 KOSPI는 조기대선 결과 이재명 대통령 당선, 상법개정과 자본시장 선진화, 2차 추경을 통한 확장적 재정정책 및 신성장 산업 지원, 내수부양의지 표명 등 정책 기대감 유입, 외국인 수급 개선되며 코리아 디스카운트 해소 국면. 관세로 인한 글로벌 경기 우려 완화국면 지속. 한편 중동지역 지정학적 긴장은 미국 개입 이후 이란의 제한적 보복, 이후 미국 중재로 휴전협정 성사되며 위험 선호 심리 회복. 이재명 대통령의 선거기간 주요 공약으로 '코스피 5000 시대' 등 증시 부양에 대한 강한 의지 표명, 특히 취임 2~3주 안에 상법개정안 처리를 공언하며 지주사, 금융지주, 신재생 등 밸류에이션 저평가 종목들이 정책 기대감에 강세를 보였으나, 상법 개정안은 국회 본회의 일정 지연되며 계류 중. 이재명 행정부는 19일 국무회의를 통해 30조원 규모의 2차 추가경정예산안을 편성하였으며, 내수소비 지원과 건설경기 활성화, 신산업 투자 촉진에 활용될 예정. 특히 최근 국가 AI 수석으로 민간 출신 전문가인 하정우 수석을 발탁하며 소버린 AI 등을 국가 성장 동력으로 추진. 정부의 원화 스테이블 코인 도입 기대감과 글로벌 제도화 흐름 속에서 관련주가 폭등하면서 현재 주식시장의 핵심 테마로 부상. 6월 중순 이후 한국 증시는 주도주 조정 속 소외주 순환매가 두드러진 가운데 KOSPI 지수는 3,071.70pt로 마감. 스타일은 대형주 비중이 높은 KOSPI200 지수는 414.60pt로 +15.29% 상승하였고 중소형주 비중이 높은 KOSDAQ 지수는 +6.42% 상승하며 마감.
""",
    },

    "(대신)트러스톤-2025년 8월 주식 전망.docx": {
        "target_month": "2025년 7월",
        "next_month"  : "2025년 8월",
        "report": """[7월 리뷰]
1.	전월동향: KOSPI 3,245.44pt(+5.66%)
- 7월 KOSPI는 7월 9일로 예정되어있던 상호관세 시한이 8월로 연장되어, 관세 불확실성 완화되며 글로벌 증시 상승 지속. 7월 31일 트럼프 미국 대통령이 상호관세율을 기존 25%에서 15%로 낮추는 등의 내용을 담은 행정명령에 서명하면서 한국 경제의 성장률 리스크는 대부분 해소되었으며, 관세 영향은 제한적. 국내 증시는 상법개정 이후 저평가 요인인 거버넌스 해소를 위한 정치권 움직임 이어졌으며, 7월 3일 상법개정안 여야 합의로 국회 본회의를 통과한 데 이어 '코스피5000특위'를 중심으로 자사주 취득 시 1년 이내 소각 의무화, 집중투표제 등 관련 법안 논의 활발히 진행. 이외에도 기재부의 세법 개정이 예고되며 배당소득 분리과세 등 다양한 방안으로 증시 저평가 해소방안 논의, 정책 기대감도 유입되며, 외국인 중심으로 순매수 지속. 2분기 실적 시즌이 시작되면서 삼성전자는 2분기 매출 74조원, 영업이익 4.6조원으로 컨센서스 하회하는 어닝쇼크 기록. 실적 저점 확인한 뒤 불확실성 해소되며 주가는 상승. 엔비디아의 중국향 H20 수출 승인 소식과, 삼성전자 이재용 회장 사법 리스크 해소 등 가세. 반면, SK하이닉스는 HBM 가격경쟁 심화 전망과 함께 과열해소 국면. 이외에도 주요 업종/종목별 실적 전망과 결과 발표에 따른 키맞추기 및 순환매가 두드러진 가운데 KOSPI 지수는 3,245.44pt로 마감. 스타일은 대형주 비중이 높은 KOSPI200 지수는 438.60pt로 +5.79% 상승하였고 중소형주 비중이 높은 KOSDAQ 지수는 +3.04% 상승하며 마감.
""",
    },

    "(대신)트러스톤-2025년 9월 주식 전망.docx": {
        "target_month": "2025년 8월",
        "next_month"  : "2025년 9월",
        "report": """[8월 리뷰]
1.	전월동향: KOSPI 3,186.01pt(-1.83%)
- 8월 KOSPI는 관세협상이 마무리된 이후 협상 기대감 관련 이슈 소멸. 7월 고용보고서 부진과 8월 잭슨홀 미팅에서 고용을 근거로 금리를 내릴 수 있다고 발언하며 연준의 9월 금리인하 전망이 기정사실화. 월초 기획재정부가 발표한 세제개편안에서 배당소득 분리과세 최고세율이 기대했던 25%보다 높은 35%로 발표하며 지주가 약세를 보였고, 주식 양도소득세에 해당되는 대주주 기준이 50억원에서 10억원으로 하향된 점도 증권주에 부담으로 작용. 이로 인해 세제개편안에 대한 실망감이 커지며 KOSPI는 약 4% 하락 출발. 원자력발전 업종은 한국수력원자력과 미국 웨스팅하우스사의 계약 세부 내용이 '팀코리아' 원전 수출에 불리하게 작성되었다는 우려 부각, 웨스팅하우스사에 지급해야하는 로열티와 핵심부품 수주, 북미지역 진출 제한 소식 등이 알려지며 상반기 주도주였던 원전 업종 기대감 후퇴. 다만 해당 계약이 불확실성을 없애기 위한 전략적 양보였다는 반박도 제기. 8월 25일 한미 정상회담이 훈훈한 분위기 속에 마무리되었으며 트럼프 대통령은 기자간담회에서 조선업, 알레스카 LNG투자 등을 언급. 회담 직후 한미 비즈니스 행사에서 기업 총수들 간의 협업 계약들이 공개되면서 일부 업종에 단기 모멘텀 형성. 코스피는 3,200선을 중심으로 순환매와 매물소화 과정을 거치며 주도주 차익실현 압력 강해지며 KOSPI 지수는 3,186.01pt로 하락 마감. 스타일은 대형주 비중이 높은 KOSPI200 지수는 430.12pt로 -1.93% 하락하였고 중소형주 비중이 높은 KOSDAQ 지수는 -1.03% 하락하며 마감.
""",
    },

    "(대신)트러스톤-2025년 10월 주식 전망.docx": {
        "target_month": "2025년 9월",
        "next_month"  : "2025년 10월",
        "report": """[9월 리뷰]
1.	전월동향: KOSPI 3,424.60pt(+7.49%)
- 9월 KOSPI는 미국 연준이 9개월만에 25bp 기준금리를 인하하면서 통화정책 기대감이 부각된 가운데, 반도체 업황 개선 기대감이 유입되며 큰 폭으로 상승. 특히 AI 반도체 사이클로 촉발된 DRAM 등 공급 부족 전망이 시장의 핵심 모멘텀으로 작용하면서, 삼성전자를 중심으로 한 외국인 순매수가 지수 상승을 주도. 여기에 대주주 양도세 10억원 기준 강화안을 철회하고 현행 50억원 유지 방침을 밝히면서 시장친화적 정책 의지가 확인되었고, 자본시장 부양에 대한 신뢰가 회복되면서 코스피 아웃포펌 동력으로 작용. 이재명 대통령의 취임 100일 기자회견 진행, 자본시장 정책과 산업정책(친환경에너지, K-문화) 등 정책 방향성 재차 강조. 다만 AI 버블 우려와 금리인하 기대 정점 통과, 한국은 미국과의 통상협상이 갈등 양상을 보이면서 불안심리 유입되었고 트럼프의 의약품/반도체 등 품목관세 또한 불확실성 요인으로 작용. 시장은 협상이 교착 상태로 장기화되며 불확실성이 확대되거나, 미국의 요구를 그대로 수용하는 최악의 시나리오를 선반영하며 변동성이 커짐. 단기 고평가 논란 속 차익실현 압력도 커지면서 KOSPI 지수는 3,424.60pt로 마감. 스타일은 대형주 비중이 높은 KOSPI200 지수는 474.03pt로 +10.21% 상승하였고 중소형주 비중이 높은 KOSDAQ 지수는 +5.56% 상승하며 마감.
""",
    },
    "(대신) 트러스톤-2026년 2월 주식 전망.docx": {
        "target_month": "2026년 1월",
        "next_month"  : "2026년 2월",
        "report": """[1월 리뷰]
1.	전월동향: KOSPI 5,224.36pt(+23.97%)
- 1월 KOSPI는 글로벌 증시 변동성에도 불구하고 연일 상승하며 신고가 랠리로 출발. NVIDIA와 AMD의 GPU를 시작으로 DRAM까지 글로벌 반도체 가격이 급등하며 공급 부족 인식이 강화되었고, 이에 따라 반도체 업황 기대가 재차 상향 조정되며 슈퍼사이클 전망 지속. 이러한 환경속에서 삼성전자는 2025년 4분기 잠정실적으로 영업이익 20조원을 기록하며 컨센서스를 상회, 반도체 슈퍼사이클 논리가 실적으로 확인. 반도체 산업을 중심으로 12개월 선행 EPS 상향 조정이 이어지며 지수 상승을 주도. 한편 CES 이후 시장의 시선은 피지컬 AI로 확장. 보스턴다이내믹스의 휴머노이드 로봇 ‘아틀라스’ 시연 및 양산 계획 공개와 함께 현대차그룹의 AI 로보틱스·자율주행 전략이 구체화되며 로보틱스 산업의 상업화 기대가 부각됨. 이는 그간 저평가되었던 시클리컬 업종 전반의 Re-Rating 국면 진입으로 이어지며 수급 확산을 동반. 외교적으로는 이재명–시진핑 정상회담을 계기로 한중 관계 개선 기대가 부각되었으나, 한한령 해제와 관련한 구체적 실행 계획 부재로 정책 모멘텀은 점진적 반영 국면에 머무름. 동시에 정치권에서 KOSPI 5,000 달성 이후 KOSDAQ 3,000 목표가 제시되며 투자심리가 개선, 코스닥 및 중소형주 중심의 수익률 키맞추기 흐름이 강화되면서 KOSPI 지수는 5,224.36pt로 마감. 스타일은 대형주 비중이 높은 KOSPI200 지수는 768.41pt로 +26.80% 상승하였고 중소형주 비중이 높은 KOSDAQ 지수는 +24.20% 상승하며 마감.
""",
    }
}

# 워드 처리

def extract_review_table(docx_path: str) -> str:
    doc = Document(docx_path)

    if not doc.tables:
        raise ValueError(f"테이블이 없습니다: {docx_path}")

    table = doc.tables[0]
    seen = set()
    parts = []
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text and text not in seen:
                seen.add(text)
                parts.append(text)

    result = "\n".join(parts)

    if len(result) >= 100:
        return result

    for table in doc.tables:
        seen = set()
        parts = []
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in seen:
                    seen.add(text)
                    parts.append(text)
        candidate = "\n".join(parts)
        if "리뷰" in candidate and len(candidate) >= 100:
            return candidate

    raise ValueError(f"유효한 리뷰 테이블을 찾을 수 없습니다: {docx_path}")


def generate_report(test_docx_path: str) -> str:
    messages = [{"role": "system", "content": SYSTEM_PROMPT}]

    # Training 예시 로드
    loaded = 0
    for docx_filename, example in EXAMPLE_REPORTS.items():
        docx_path = os.path.join(TRAINING_DOCX_DIR, docx_filename)
        if not os.path.exists(docx_path):
            print(f"  ⚠️  Training 파일 없음 (건너뜀): {docx_path}")
            continue
        try:
            input_text = extract_review_table(docx_path)
        except Exception as e:
            print(f"  ⚠️  텍스트 추출 실패: {docx_path}\n     → {e}")
            continue

        messages.append({
            "role": "user",
            "content": (
                "다음 주식시장 월간 리뷰 원문을 읽고, 정해진 형식에 맞춰 요약 보고서를 작성하세요.\n\n"
                f"[원문]\n{input_text[:8000]}"
            )
        })
        messages.append({
            "role": "assistant",
            "content": example["report"]  # ← 수정된 부분
        })
        loaded += 1
        print(f"  ✅ Training 예시 로드: {docx_filename}")

    print(f"  총 {loaded}개 예시 로드 완료")

    print(f"\n  📄 Test 파일 추출 중: {os.path.basename(test_docx_path)}")
    input_text = extract_review_table(test_docx_path)
    print(f"  추출 완료 ({len(input_text):,}자)")

    messages.append({
        "role": "user",
        "content": (
            "다음 주식시장 월간 리뷰 원문을 읽고, 정해진 형식에 맞춰 요약 보고서를 작성하세요.\n\n"
            f"[원문]\n{input_text[:8000]}"
        )
    })

    #  GPT 호출
    print(f"  🤖 GPT 호출 중 (model={MODEL})...")
    response = client.chat.completions.create(
        model=MODEL,
        messages=messages,
        temperature=TEMPERATURE,
        max_tokens=MAX_TOKENS,
    )

    report = response.choices[0].message.content
    usage  = response.usage
    print(
        f"  토큰 사용: input={usage.prompt_tokens:,} / "
        f"output={usage.completion_tokens:,} / "
        f"total={usage.total_tokens:,}"
    )
    return report



def main():
    print("=" * 60)
    print("  한국 주식시장 월간 리뷰 요약 보고서 생성기")
    print(f"  Test 파일: {os.path.basename(TEST_DOCX_PATH)}")
    print("=" * 60)

    if OPENAI_API_KEY.startswith("sk-proj-xxx"):
        print("\n⚠️  OPENAI_API_KEY를 설정하세요 (코드 상단 OPENAI_API_KEY 변수)")
        return None

    if not os.path.exists(TEST_DOCX_PATH):
        print(f"\n❌ Test 파일을 찾을 수 없습니다: {TEST_DOCX_PATH}")
        return None

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("\n" + "=" * 60)
    print("  [Step 1] 요약 보고서 생성")
    print("=" * 60)

    try:
        report = generate_report(TEST_DOCX_PATH)
    except Exception as e:
        print(f"\n❌ 보고서 생성 실패: {e}")
        return None

    print("\n" + "=" * 60)
    print("  [생성된 보고서]")
    print("=" * 60)
    print(report)

    base_name   = os.path.splitext(os.path.basename(TEST_DOCX_PATH))[0]
    output_name = base_name + "_요약보고서.txt"
    output_path = os.path.join(OUTPUT_DIR, output_name)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(report)

    print(f"\n✅ 저장 완료: {output_path}")
    print("=" * 60)
    return report


if __name__ == "__main__":
    main()

# python BOT/api_리뷰.py